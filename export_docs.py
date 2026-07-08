import docx

doc = docx.Document()
with open('/Users/dima/.gemini/antigravity-ide/brain/69a66d33-ff53-4aa1-8c87-2f6395420a88/TZ_FOR_IT.md', 'r') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if line.startswith('# '):
        doc.add_heading(line[2:], 0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], 1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], 2)
    elif line.startswith('* '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
        doc.add_paragraph(line, style='List Number')
    elif line.startswith('```'):
        continue
    else:
        if line:
            doc.add_paragraph(line)

doc.save('/Users/dima/Desktop/hello-park-invite/TZ_FOR_IT.docx')
print("Successfully generated TZ_FOR_IT.docx")
